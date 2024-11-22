"""This script generates a PDF file for each record in a excel"""

try:
    import locale
    import os
    import re
    import sys
    import traceback
    from datetime import datetime

    import pandas as pd
    import pyfiglet
    from docx import Document
    from docx2pdf import convert
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Inches

    FILES_PATH = os.path.dirname(os.path.abspath(__file__)) + "\\"

    def disclaimer():
        ascii_art = pyfiglet.figlet_format("C & C", font="doh")

        lines = ascii_art.splitlines()

        # Find the first non-empty line from the top
        start_line = 0
        for i, line in enumerate(lines):
            if line.strip():
                start_line = i
                break

        # Find the last non-empty line from the bottom
        end_line = len(lines) - 1
        for i in range(len(lines) - 1, -1, -1):
            if lines[i].strip():
                end_line = i
                break

        # Crop the ASCII art to remove empty lines
        ASCCI = "\n".join(lines[start_line : end_line + 1])

        print(ASCCI)
        print()

        print(
            "\nRecuerde que este archivo es privado y no debe ser compartido con personal no autorizado.\n"
        )

        print("RECUERDE CERRAR MICROSOFT WORD ANTES DE EJECUTAR ESTE SCRIPT.\n")
        input("Presiona enter para continuar.")

    # Check if the folders exists and create them if they don't
    os.makedirs(FILES_PATH + "pdf", exist_ok=True)

    if not os.path.exists(FILES_PATH + "Plantillas"):
        print("No se encontró la carpeta 'Plantillas'.")
        input("Presiona enter para salir.")
        sys.exit()

    if not os.path.exists(FILES_PATH + "images"):
        print("No se encontró la carpeta 'images'.")
        input("Presiona enter para salir.")
        sys.exit()

    if not os.path.exists(FILES_PATH + "images/logo.png"):
        print("No se encontró el archivo 'logo.png'.")
        input("Presiona enter para salir.")
        sys.exit()

    if not os.path.exists(FILES_PATH + "BASE BOT CORRESPONDENCIA.xlsx"):
        print("No se encontró el archivo 'BASE BOT CORRESPONDENCIA.xlsx'.")
        input("Presiona enter para salir.")
        sys.exit()

    def add_image_as_header(doc_path, img_path):
        # Load the document
        doc = Document(doc_path)

        # Access the header of the first section
        section = doc.sections[0]
        header = section.header

        # Add a paragraph to the right of the header at the right
        paragraph = header.paragraphs[0]
        paragraph.alignment = (
            WD_PARAGRAPH_ALIGNMENT.RIGHT
        )  # Align the paragraph to the right
        run = paragraph.add_run()
        run.add_picture(img_path, width=Inches(0.6))

        # Save the modified document
        temp_path = os.path.join(os.path.dirname(doc_path), "temp_with_header.docx")
        doc.save(temp_path)
        return temp_path

    def get_data_from_excel(file):
        """Get the data from the excel file."""
        try:
            # Read the data from the excel file
            data = pd.read_excel(file)
            # Ensure that the columns are in the correct order
            required_columns = [
                "CEDULA",
                "CODIGO",
                "DIRECTORA",
                "NOMBRE DE DIRECTORA",
                "CORREO DIR",
                "NOMBRE CONSULTORA",
                "DIRECCION",
                "BARRIO",
                "CIUDAD",
                "FACTURA",
                "VENCIMIENTO",
                "VALOR TOTAL AL DIA",
                "DIAS MORA AL DIA",
                "EDAD DE LIQUIDACION",
            ]
            for column in required_columns:
                if column not in data.columns:
                    print(
                        f"La columna '{column}' no se encuentra en el archivo '{file}'."
                    )
                    input("Presiona enter para salir.")
                    sys.exit()
            return data
        except Exception as e:
            print(f"Hubo un error al leer el archivo {file}: {str(e)}")
            input("Presiona enter para salir.")
            sys.exit()

    # Function to load and replace placeholders in the .docx
    # def load_and_replace_docx(docx_path, replacements):
    #     doc = Document(docx_path)
    #     # Replace placeholders in the body of the document
    #     for paragraph in doc.paragraphs:
    #         for placeholder, new_value in replacements.items():
    #             if placeholder in paragraph.text:
    #                 # Replace placeholder text while keeping formatting
    #                 paragraph.text = paragraph.text.replace(placeholder, new_value)

    #     temp_docx_path = FILES_PATH + "words/modified.docx"
    #     doc.save(temp_docx_path)
    #     return temp_docx_path

    def load_and_replace_docx(docx_path, replacements):
        doc = Document(docx_path)
        for paragraph in doc.paragraphs:
            for placeholder, new_value in replacements.items():
                if placeholder in paragraph.text:
                    # Replace placeholder with actual value
                    paragraph.text = re.sub(placeholder, new_value, paragraph.text)

        # Save modified document to a temporary file
        modified_docx_path = FILES_PATH + "words/modified.docx"
        doc.save(modified_docx_path)
        return modified_docx_path

    def select_template(age: str):
        if age.startswith("001") or age.startswith("010"):
            original_docx_path = FILES_PATH + "Plantillas/Plantilla 10-30.docx"
        elif age.startswith("031"):
            original_docx_path = FILES_PATH + "Plantillas/Plantilla 31-60.docx"
        elif age.startswith("061"):
            original_docx_path = FILES_PATH + "Plantillas/Plantilla 61-120.docx"
        elif age.startswith("121"):
            original_docx_path = FILES_PATH + "Plantillas/Plantilla 121-180.docx"
        elif age.startswith("181") or age.startswith("271"):
            original_docx_path = FILES_PATH + "Plantillas/Plantilla 180+.docx"
        elif age.startswith("CASTIGO"):
            original_docx_path = FILES_PATH + "Plantillas/Plantilla Castigo.docx"
        else:
            print(f"Plantilla de liquidación no encontrada para {data['CEDULA'][i]}")
            input("Presiona enter para salir.")
            sys.exit()
        return original_docx_path

    if __name__ == "__main__":
        disclaimer()
        TODAY = datetime.now().strftime("%d/%m/%Y")
        locale.setlocale(locale.LC_ALL, "es_CO.UTF-8")
        # Paths to the files
        print("Leyendo 'BASE BOT CORRESPONDENCIA.xlsx'...")
        data = get_data_from_excel(FILES_PATH + "BASE BOT CORRESPONDENCIA.xlsx")
        print("Creando PDFs...")
        errors = []
        for i in range(len(data["CEDULA"])):
            original_docx_path = select_template(str(data["EDAD DE LIQUIDACION"][i]))
            os.makedirs(FILES_PATH + "pdf/" + str(data["DIRECTORA"][i]), exist_ok=True)
            output_pdf_path = (
                FILES_PATH + f"/pdf/{str(data['DIRECTORA'][i])}/{data['CEDULA'][i]}.pdf"
            )
            # Dictionary with placeholders and their corresponding values
            replacements = {
                "XDATE_TODAYX": TODAY,
                "XCONSULTANT_NAMEX": data["NOMBRE CONSULTORA"][i],
                "XADDRESSX": data["DIRECCION"][i],
                "XCITYX": data["CIUDAD"][i],
                "XBILLX": str(data["FACTURA"][i]),
                "XEXPIRATION_DATEX": data["VENCIMIENTO"][i].strftime("%d/%m/%Y"),
                # format like money
                "XVALUEX": locale.currency(
                    data["VALOR TOTAL AL DIA"][i], grouping=True
                ).split(",")[0],
                "XDIASX": str(data["DIAS MORA AL DIA"][i]),
            }
            # Load, replace placeholders, and save as a new .docx file
            modified_docx_path = load_and_replace_docx(original_docx_path, replacements)
            # Add an image to the top right of the document
            img_path = FILES_PATH + "images/logo.png"
            modified_docx_path = add_image_as_header(modified_docx_path, img_path)

            try:
                # Convert the modified .docx to PDF
                convert(modified_docx_path, output_pdf_path)
                print("PDF created successfully at:", output_pdf_path)
            except Exception as e:
                print(f"Error creating PDF for {data['CEDULA'][i]}: {str(e)}")
                print("Volviendo a intentar...")
                try:
                    convert(modified_docx_path, output_pdf_path)
                    print("PDF created successfully at:", output_pdf_path)
                except Exception as e:
                    print(f"Error creating PDF for {data['CEDULA'][i]}: {str(e)}")
                    errors.append(data["CEDULA"][i])
        if errors:
            print("Hubo un error al crear los siguientes PDFs")
            for error in errors:
                print(error)
        input("Presiona enter para salir.")

except Exception as e:
    print(traceback.format_exc())
    print(f"Hubo un error: {str(e)}")
    input("Presiona enter para salir.")
    sys.exit()

finally:
    os.system("taskkill /f /im WINWORD.EXE")
